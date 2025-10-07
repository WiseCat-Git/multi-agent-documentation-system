import streamlit as st
import os
from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
from main import generate_documentation
import logging
from dotenv import load_dotenv
import time

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Function to create a .docx file with formatting
def create_docx(text, filename):
    document = Document()
    lines = text.split("\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("###"):
            document.add_heading(line.replace("###", "").strip(), level=3)
        elif line.startswith("##"):
            document.add_heading(line.replace("##", "").strip(), level=2)
        elif line.startswith("#"):
            document.add_heading(line.replace("#", "").strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            document.add_paragraph(line.strip(), style="List Number")
        else:
            document.add_paragraph(line)
    
    document_path = f"{filename}.docx"
    document.save(document_path)
    return document_path

# Function to read file content based on file type
def read_file(uploaded_file):
    try:
        if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            file_content = "\n".join([para.text for para in doc.paragraphs])
        elif uploaded_file.type == "application/pdf":
            pdf = PdfReader(uploaded_file)
            file_content = ""
            for page in pdf.pages:
                file_content += page.extract_text()
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            df = pd.read_excel(uploaded_file)
            file_content = df.to_string()
        elif uploaded_file.type == "text/csv":
            df = pd.read_csv(uploaded_file)
            file_content = df.to_string()
        else:
            file_content = None
        return file_content
    except Exception as e:
        logger.error(f"Error reading file: {e}")
        return None

def main():
    st.set_page_config(
        page_title="AI Multi-Agent Documentation System",
        page_icon="🤖",
        layout="wide"
    )
    
    st.title("🤖 AI Multi-Agent Documentation System")
    st.markdown("""
    ### Powered by CrewAI
    This system uses **4 specialized AI agents** working together:
    - 🔍 **Researcher**: Gathers and analyzes information
    - 📝 **Documentation Specialist**: Creates structured documentation
    - ✅ **Quality Reviewer**: Ensures accuracy and completeness
    - 🎓 **Training Specialist**: Converts docs into learning materials
    """)
    
    # CSS for styling
    st.markdown("""
    <style>
    .stAlert {
        padding: 1rem;
        margin: 1rem 0;
    }
    .success-box {
        padding: 1rem;
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        margin: 1rem 0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Sidebar configuration
    st.sidebar.title("⚙️ Configuration")
    
    # Input method selection
    method = st.sidebar.radio(
        "Choose Input Method",
        ("Text Input", "File Upload"),
        help="Select how you want to provide content"
    )
    
    # Template selection
    template = st.sidebar.selectbox(
        "Choose Documentation Template",
        [
            "Technical Documentation",
            "Meeting Minutes",
            "Standard Report",
            "API Documentation",
            "User Guide",
            "Project Proposal"
        ],
        help="Select the type of documentation you want to generate"
    )
    
    # Topic/Title input
    topic = st.sidebar.text_input(
        "Document Title/Topic",
        placeholder="e.g., 'Python Best Practices' or 'Q4 Planning Meeting'",
        help="Provide a clear title or topic for your documentation"
    )
    
    text_input = ""
    file_content = ""

    # Content input area
    st.sidebar.markdown("---")
    st.sidebar.subheader("📄 Content Input")
    
    if method == "Text Input":
        text_input = st.sidebar.text_area(
            "Enter your content",
            placeholder="Describe what you want documented...",
            height=250,
            help="Provide detailed information about what you want documented"
        )
    elif method == "File Upload":
        uploaded_file = st.sidebar.file_uploader(
            "Upload your file",
            type=["docx", "pdf", "xlsx", "csv"],
            help="Upload a document to extract content from"
        )
        if uploaded_file:
            with st.spinner("Reading file..."):
                file_content = read_file(uploaded_file)
                if file_content is None:
                    st.sidebar.error("❌ Unsupported file type or error reading file.")
                else:
                    st.sidebar.success(f"✅ File loaded: {uploaded_file.name}")
                    with st.sidebar.expander("Preview file content"):
                        st.text(file_content[:500] + "..." if len(file_content) > 500 else file_content)

    # API Key validation
    st.sidebar.markdown("---")
    openai_key = os.getenv('OPENAI_API_KEY')
    serper_key = os.getenv('SERPER_API_KEY')
    
    if not openai_key:
        st.sidebar.error("❌ OPENAI_API_KEY not found in environment")
    else:
        st.sidebar.success("✅ OpenAI API Key loaded")
    
    if not serper_key:
        st.sidebar.warning("⚠️ SERPER_API_KEY not found (optional for research)")
    else:
        st.sidebar.success("✅ Serper API Key loaded")

    # Main action buttons
    col1, col2 = st.columns(2)
    
    with col1:
        generate_button = st.button(
            "🚀 Generate Documentation",
            type="primary",
            use_container_width=True,
            disabled=not topic or (not text_input and not file_content)
        )
    
    with col2:
        preview_button = st.button(
            "👁️ Preview Input",
            use_container_width=True,
            disabled=not topic or (not text_input and not file_content)
        )

    # Preview functionality
    if preview_button:
        st.subheader("📋 Input Preview")
        st.write(f"**Topic:** {topic}")
        st.write(f"**Template:** {template}")
        st.write(f"**Input Method:** {method}")
        
        content = text_input if text_input else file_content
        with st.expander("View Content", expanded=True):
            st.text_area("Content", content, height=300, disabled=True)

    # Generate documentation
    if generate_button:
        if not openai_key:
            st.error("❌ Cannot generate documentation without OpenAI API key. Please set OPENAI_API_KEY in your .env file.")
            return
        
        content = text_input if text_input else file_content
        
        # Progress indicators
        st.markdown("---")
        st.subheader("🔄 Multi-Agent Processing Pipeline")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Agents progress tracking
        agents_status = st.container()
        with agents_status:
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                researcher_status = st.empty()
                researcher_status.info("🔍 Researcher: Waiting...")
            with col2:
                doc_status = st.empty()
                doc_status.info("📝 Doc Specialist: Waiting...")
            with col3:
                reviewer_status = st.empty()
                reviewer_status.info("✅ Reviewer: Waiting...")
            with col4:
                trainer_status = st.empty()
                trainer_status.info("🎓 Trainer: Waiting...")
        
        try:
            # Simulate agent progress updates
            researcher_status.warning("🔍 Researcher: Working...")
            progress_bar.progress(10)
            status_text.text("Phase 1/4: Researching and analyzing topic...")
            
            # Call the CrewAI system
            start_time = time.time()
            result = generate_documentation(topic, template, content)
            end_time = time.time()
            
            # Update progress
            researcher_status.success("🔍 Researcher: Complete ✓")
            progress_bar.progress(40)
            
            doc_status.success("📝 Doc Specialist: Complete ✓")
            progress_bar.progress(70)
            
            reviewer_status.success("✅ Reviewer: Complete ✓")
            progress_bar.progress(90)
            
            trainer_status.success("🎓 Trainer: Complete ✓")
            progress_bar.progress(100)
            
            status_text.text(f"✅ Complete! Generated in {end_time - start_time:.2f} seconds")
            
            # Store in session state
            st.session_state["result"] = result
            st.session_state["topic"] = topic
            
            # Success message
            st.markdown("---")
            st.success("🎉 Documentation generated successfully!")
            
            # Create downloadable file
            try:
                docx_path = create_docx(str(result), "Documentation")
                
                with open(docx_path, "rb") as f:
                    st.download_button(
                        label="📥 Download as DOCX",
                        data=f,
                        file_name=f"{topic.replace(' ', '_')}_Documentation.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            except Exception as e:
                logger.error(f"Error creating DOCX: {e}")
                st.warning("Could not create DOCX file, but you can still copy the content below.")
            
            # Display result
            st.markdown("---")
            st.subheader("📄 Generated Documentation")
            
            # Show both versions
            tab1, tab2 = st.tabs(["Formatted View", "Raw Content"])
            
            with tab1:
                st.markdown(str(result))
            
            with tab2:
                st.code(str(result), language="markdown")
                
        except Exception as e:
            logger.error(f"Error during documentation generation: {e}")
            st.error(f"❌ Error: {str(e)}")
            st.info("💡 Tip: Make sure all required API keys are set in your .env file")
            return

    # Show previous result if exists
    if "result" in st.session_state and not generate_button:
        st.markdown("---")
        st.subheader("📄 Last Generated Documentation")
        st.info(f"Topic: {st.session_state.get('topic', 'N/A')}")
        
        with st.expander("View Previous Result", expanded=False):
            st.markdown(str(st.session_state["result"]))

if __name__ == "__main__":
    main()